import base64
import json
import re
import copy
from typing import Dict, List, Any
import requests
import streamlit as st
import io
import pandas as pd
import hashlib

#report imports
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem
from xml.sax.saxutils import escape

#st.set_option("global.dataFrameSerialization", "legacy")

#FastAPI connect
BACKEND_BASE_URL = st.secrets["BACKEND_BASE_URL"].rstrip("/")
BACKEND_GRADE_URL = f"{BACKEND_BASE_URL}/grade"
BACKEND_SUMMARY_URL = f"{BACKEND_BASE_URL}/summarize_batch"
BACKEND_NORM_URL = f"{BACKEND_BASE_URL}/norm_reference_batch"
BACKEND_STUDENT_REPORT_URL = f"{BACKEND_BASE_URL}/student_reports_batch"
BACKEND_TRANSCRIBE_URL = f"{BACKEND_BASE_URL}/transcribe_handwriting"
BACKEND_CHAT_URL = f"{BACKEND_BASE_URL}/voxlm_chat"
BACKEND_REFINE_MODEL_ANSWER_URL = f"{BACKEND_BASE_URL}/refine_model_answer"
BACKEND_VIDEO_MCQ_URL = f"{BACKEND_BASE_URL}/generate/questions_from_videos"
BACKEND_API_KEY = st.secrets["BACKEND_API_KEY"]

#df sanitizer
def sanitize_df_for_streamlit(df: pd.DataFrame) -> pd.DataFrame:
    clean = pd.DataFrame(df.to_dict(orient="records"))
    clean = clean.astype(object)
    clean = clean.where(pd.notnull(clean), None)
    clean.columns = [str(c) for c in clean.columns]
    return clean

#Batch comparison front helper

def reset_student_state():
    st.session_state.grade_result = None
    st.session_state.original_grade_result = None
    st.session_state.challenge_mode = False
    st.session_state.challenge_used = False
    st.session_state.challenge_submitted = False
    st.session_state.last_grade_payload = None
    st.session_state.debug_prompt_view = ""
    st.session_state.last_has_subquestions = False
    st.session_state.teacher_finalised = False
    st.session_state.finalised_grade_result = None
    st.session_state.voxlm_chat_open = False
    st.session_state.voxlm_chat_history = []
    st.session_state.voxlm_chat_input = ""


    st.session_state.pop("challenge_reason_text", None)
    st.session_state.pop("student_answer", None)

    st.session_state.pop("handwritten_response_mode", None)
    st.session_state.pop("handwritten_response_image", None)
    st.session_state.pop("last_handwritten_upload_sig", None)
    st.session_state.pop("handwriting_transcription_confidence", None)
    st.session_state.pop("handwriting_transcription_debug", None)

def start_challenge():
    st.session_state.challenge_mode = True
    st.session_state.challenge_used = True

def reset_summary_state():
    st.session_state.batch_summary_result = None
    st.session_state.batch_summary_debug_prompt = ""

def reset_student_reports_state():
    st.session_state.student_reports_result = None
    st.session_state.student_reports_debug_prompt = ""

def reset_video_mcq_state():
    st.session_state.video_mcq_result = None
    st.session_state.video_mcq_debug_prompt = ""


#Three panel Vox-LM Prototype

#Highlighter for mid panel
def build_highlighted_html(
    text: str,
    correct_segments: List[str],
    out_of_scope_segments: List[str],
    misconception_segments: List[str],
    uncertain_segments: List[str],
) -> str:
    """
    Segment-based highlighter.

    Green  = Correct
    Yellow = Out of scope
    Orange = Misconception
    Blue   = Uncertain

    The model should return exact short phrases from the student response.
    This function searches those phrases in the original response and colours them.
    It avoids nested overlapping highlights.
    """
    import html

    text = str(text or "")

    colour_map = {
        "correct": {
            "colour": "#c8f7c5", #light green
            "label": "Correct",
            "priority": 1,
        },
        "out_of_scope": {
            "colour": "#fff5c5", #light yellow
            "label": "Out of scope",
            "priority": 2,
        },
        "misconception": {
            "colour": "#ffd6a5", #light orange
            "label": "Misconception",
            "priority": 3,
        },
        "uncertain": {
            "colour": "#cfe8ff", #light blue
            "label": "Uncertain",
            "priority": 4,
        },
    }

    raw_segments = []

    def add_segments(category: str, segments: List[str]):
        seen = set()
        for seg in segments or []:
            seg = str(seg or "").strip()
            if not seg:
                continue
            key = seg.lower()
            if key in seen:
                continue
            seen.add(key)

            for m in re.finditer(re.escape(seg), text, flags=re.I):
                raw_segments.append({
                    "start": m.start(),
                    "end": m.end(),
                    "category": category,
                    "matched_text": text[m.start():m.end()],
                    "priority": colour_map[category]["priority"],
                })

    add_segments("correct", correct_segments)
    add_segments("out_of_scope", out_of_scope_segments)
    add_segments("misconception", misconception_segments)
    add_segments("uncertain", uncertain_segments)

    raw_segments.sort(
        key=lambda x: (
            x["start"],
            -(x["end"] - x["start"]),
            x["priority"],
        )
    )

    accepted = []
    occupied = [False] * len(text)

    for seg in raw_segments:
        if any(occupied[i] for i in range(seg["start"], seg["end"])):
            continue
        for i in range(seg["start"], seg["end"]):
            occupied[i] = True
        accepted.append(seg)

    accepted.sort(key=lambda x: x["start"])

    pieces = []
    pos = 0

    for seg in accepted:
        if seg["start"] > pos:
            pieces.append(html.escape(text[pos:seg["start"]]))

        category = seg["category"]
        colour = colour_map[category]["colour"]
        label = colour_map[category]["label"]
        segment_text = html.escape(text[seg["start"]:seg["end"]])

        pieces.append(
            f'<span title="{label}" '
            f'style="background-color:{colour}; padding:1px 3px; border-radius:3px;">'
            f'{segment_text}</span>'
        )
        pos = seg["end"]

    if pos < len(text):
        pieces.append(html.escape(text[pos:]))

    return "".join(pieces)



#parsers for text box
def parse_subquestions_from_text(text: str) -> List[Dict]:
    """
    Parse subquestions from free text.

    Expected structure per subquestion block (separated by blank lines):

    1:
    Question text
    Sub-score for question
    Rubric / model answer (optional, can span multiple lines)

    First line may be of the form: "1", "1:", "1.", or "1)".
    The ID is normalised to the digits only (e.g. "1").
    """
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    subqs: List[Dict] = []

    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if len(lines) < 3:
            continue

        first = lines[0]
        # Normalise ID: accept "1", "1:", "1.", "1)"
        m = re.match(r"^(\d+)\s*[:\.\)\-]?\s*$", first)
        if m:
            sid = m.group(1)  # just digits, e.g. "1"
        else:
            sid = first.strip()

        prompt = lines[1]

        try:
            max_score = float(lines[2])
        except Exception:
            max_score = None

        rubric = ""
        if len(lines) > 3:
            rubric = "\n".join(lines[3:])

        subqs.append(
            {
                "id": sid,
                "prompt": prompt,
                "max_score": max_score,
                "rubric": rubric,
            }
        )

    return subqs


def parse_few_shot_from_text(
    text: str,
    has_subquestions: bool,
    subquestions: List[Dict],
) -> List[Dict]:
    """
    Parse few-shot examples from free text.

    Expected structure per example block (separated by blank lines):

    Student Number: S1
    Answer:
    <student's full answer, possibly multiple lines>
    Marker score: 4.5

    If subquestions exist, the same answer text is used
    for every subquestion id.
    If no subquestions, stored under key "overall".
    """
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    few_shot: List[Dict] = []

    for block in blocks:
        lines = block.splitlines()
        student_id = None
        answer_lines: List[str] = []
        score = None

        idx_answer_start = None
        idx_marker = None

        for i, line in enumerate(lines):
            l = line.strip()
            lower = l.lower()
            if lower.startswith("student number:"):
                student_id = l.split(":", 1)[1].strip() or f"student_{len(few_shot) + 1}"
            elif lower.startswith("answer:"):
                idx_answer_start = i + 1
            elif lower.startswith("marker score:"):
                idx_marker = i
                try:
                    score_str = l.split(":", 1)[1].strip()
                    score = float(score_str)
                except Exception:
                    score = 0.0

        if idx_answer_start is not None:
            if idx_marker is None:
                answer_lines = [l.strip() for l in lines[idx_answer_start:]]
            else:
                answer_lines = [l.strip() for l in lines[idx_answer_start:idx_marker]]

        answer_text = "\n".join(answer_lines).strip()

        if not student_id and not answer_text:
            continue

        if score is None:
            score = 0.0

        if has_subquestions and subquestions:
            answers = {str(sq.get("id")): answer_text for sq in subquestions}
        else:
            answers = {"overall": answer_text}

        few_shot.append(
            {
                "response_id": student_id or f"student_{len(few_shot) + 1}",
                "answers": answers,
                "marker_score": score,
            }
        )

    return few_shot


def parse_student_answers_from_text(
    text: str,
    subquestions: List[Dict],
) -> Dict[str, str]:
    """
    Parse student answers into a dict keyed by subquestion id.

    Uses the IDs defined in `subquestions` and supports answer formats like:

    1:
    The length is around 9mm

    2. I will use a molar size pontic for 36 ...

    3) As the pontic will be molar sized ...

    4 The path of insertion will be tilted mesially and palatally.

    Blank lines are allowed but not required. If a subquestion
    receives no text, its value will be an empty string.
    """
    valid_ids = {str(sq.get("id")) for sq in subquestions}

    # Accumulate lines per subquestion
    answers_lines: Dict[str, List[str]] = {sid: [] for sid in valid_ids}
    current_sid: str | None = None

    # Header regex:
    #  "1:"      -> id="1", rest=""
    #  "1. ans"  -> id="1", rest="ans"
    #  "2) ans"  -> id="2", rest="ans"
    #  "3 ans"   -> id="3", rest="ans"
    header_re = re.compile(r"^(\d+)\s*[:\.\)\-]?\s*(.*)$")

    for raw in text.splitlines():
        line = raw.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            if current_sid is not None:
                answers_lines[current_sid].append("")
            continue

        m = header_re.match(stripped)
        if m:
            sid_candidate, rest = m.group(1), m.group(2)
            if sid_candidate in valid_ids:
                current_sid = sid_candidate
                if rest:
                    answers_lines[current_sid].append(rest)
                continue

        if current_sid is not None:
            answers_lines[current_sid].append(line)

    answers: Dict[str, str] = {}
    for sq in subquestions:
        sid = str(sq.get("id"))
        joined = "\n".join(answers_lines.get(sid, [])).strip()
        answers[sid] = joined

    return answers

#misconception and oos helpers

def render_pattern_table(title: str, items: List[Any], empty_message: str):
    st.markdown(f"#### {title}")

    if not items:
        st.write(empty_message)
        return

    rows = []

    for item in items:
        if isinstance(item, dict):
            rows.append({
                "Point": str(item.get("point", "") or item.get("description", "") or "").strip(),
                "% students": item.get("percent_students", None),
                "Student count": item.get("student_count", None),
            })
        else:
            rows.append({
                "Point": str(item),
                "% students": None,
                "Student count": None,
            })

    rows = [r for r in rows if r["Point"]]

    if not rows:
        st.write(empty_message)
        return

    df = pd.DataFrame(rows)

    if "% students" in df.columns:
        df["% students"] = pd.to_numeric(df["% students"], errors="coerce")

    if "Student count" in df.columns:
        df["Student count"] = pd.to_numeric(df["Student count"], errors="coerce")

    st.dataframe(
        sanitize_df_for_streamlit(df),
        use_container_width=True,
        hide_index=True,
    )


#Report generation helpers
def safe_filename(name: str) -> str:
    name = str(name or "").strip()
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", name)
    return name[:80] or "student report"


def normalize_report_item(rep: Dict) -> Dict:
    display_name = rep.get("display_name") or rep.get("student_id") or "Student"
    overall_summary = str(rep.get("overall_summary", "") or "").strip()

    strong_areas = rep.get("strong_areas", []) or []
    weak_areas = rep.get("weak_areas", []) or []

    strong_areas = [str(x).strip() for x in strong_areas if str(x).strip()]
    weak_areas = [str(x).strip() for x in weak_areas if str(x).strip()]

    return {
        "student_id": rep.get("student_id", ""),
        "display_name": display_name,
        "overall_summary": overall_summary,
        "strong_areas": strong_areas,
        "weak_areas": weak_areas,
    }


def add_report_to_docx(doc: Document, rep: Dict):
    rep = normalize_report_item(rep)

    doc.add_heading(rep["display_name"], level=1)

    p = doc.add_paragraph()
    p.add_run("Performance summary: ").bold = True
    p.add_run(rep["overall_summary"] or "No summary available.")

    doc.add_paragraph("Strong areas", style="Heading 2")
    if rep["strong_areas"]:
        for item in rep["strong_areas"]:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("None identified.", style="List Bullet")

    doc.add_paragraph("Weak areas", style="Heading 2")
    if rep["weak_areas"]:
        for item in rep["weak_areas"]:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("None identified.", style="List Bullet")


def build_single_report_docx_bytes(rep: Dict) -> bytes:
    rep = normalize_report_item(rep)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    add_report_to_docx(doc, rep)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_all_reports_docx_bytes(reports: List[Dict], solo_analysis: Dict | None = None) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Student Test Reports", level=1)

    if solo_analysis:
        summary = str(solo_analysis.get("question_level_summary", "") or "").strip()
        if summary:
            p = doc.add_paragraph()
            p.add_run("Question summary: ").bold = True
            p.add_run(summary)

    for i, rep in enumerate(reports):
        if i > 0:
            doc.add_page_break()
        add_report_to_docx(doc, rep)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_report_to_pdf_story(story: List, rep: Dict, styles):
    rep = normalize_report_item(rep)

    heading_style = styles["Heading2"]
    body_style = styles["BodyText"]

    story.append(Paragraph(escape(rep["display_name"]), heading_style))
    story.append(Spacer(1, 4 * mm))

    summary_text = escape(rep["overall_summary"] or "No summary available.")
    story.append(Paragraph(f"<b>Performance summary:</b> {summary_text}", body_style))
    story.append(Spacer(1, 4 * mm))

    story.append(Paragraph("Strong areas", heading_style))
    if rep["strong_areas"]:
        strong_list = ListFlowable(
            [
                ListItem(Paragraph(escape(item), body_style))
                for item in rep["strong_areas"]
            ],
            bulletType="bullet",
        )
        story.append(strong_list)
    else:
        story.append(ListFlowable([ListItem(Paragraph("None identified.", body_style))], bulletType="bullet"))

    story.append(Spacer(1, 4 * mm))

    story.append(Paragraph("Weak areas", heading_style))
    if rep["weak_areas"]:
        weak_list = ListFlowable(
            [
                ListItem(Paragraph(escape(item), body_style))
                for item in rep["weak_areas"]
            ],
            bulletType="bullet",
        )
        story.append(weak_list)
    else:
        story.append(ListFlowable([ListItem(Paragraph("None identified.", body_style))], bulletType="bullet"))

    story.append(Spacer(1, 6 * mm))


def build_single_report_pdf_bytes(rep: Dict) -> bytes:
    rep = normalize_report_item(rep)

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    styles = getSampleStyleSheet()
    styles["BodyText"].fontName = "Helvetica"
    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading2"].fontName = "Helvetica-Bold"

    story = []
    add_report_to_pdf_story(story, rep, styles)
    pdf.build(story)

    buffer.seek(0)
    return buffer.getvalue()


def build_all_reports_pdf_bytes(reports: List[Dict], solo_analysis: Dict | None = None) -> bytes:
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    styles = getSampleStyleSheet()
    styles["BodyText"].fontName = "Helvetica"
    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading2"].fontName = "Helvetica-Bold"

    story = []
    story.append(Paragraph("Student Test Reports", styles["Title"]))
    story.append(Spacer(1, 6 * mm))

    if solo_analysis:
        summary = str(solo_analysis.get("question_level_summary", "") or "").strip()
        if summary:
            story.append(Paragraph(f"<b>Question summary:</b> {escape(summary)}", styles["BodyText"]))
            story.append(Spacer(1, 6 * mm))

    for i, rep in enumerate(reports):
        if i > 0:
            story.append(PageBreak())
        add_report_to_pdf_story(story, rep, styles)

    pdf.build(story)
    buffer.seek(0)
    return buffer.getvalue()

#helpers for mcq generation from video lectures
def format_seconds(value: Any) -> str:
    """
    Format seconds as MM:SS, e.g.
    270 -> 04:30
    78.94 -> 01:18

    If duration is 1 hour or more, format as H:MM:SS.
    """
    try:
        seconds = int(float(value))
    except Exception:
        return ""

    if seconds < 0:
        seconds = 0

    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60

    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"

    return f"{minutes:02d}:{secs:02d}"

def flatten_video_mcq_rows(video_mcq_result: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows = []

    if not video_mcq_result:
        return rows

    def add_question(q: Dict[str, Any]):
        if not isinstance(q, dict):
            return

        options = q.get("options", []) or []
        opt_lookup = {}

        for opt in options:
            if isinstance(opt, dict):
                label = str(opt.get("label", "")).strip()
                text = str(opt.get("text", "")).strip()
                if label:
                    opt_lookup[label] = text

        qf = q.get("quality_flags", {}) or {}

        rows.append(
            {
                "question_id": q.get("question_id", ""),
                "question_type": q.get("question_type", ""),
                "question_kind": q.get("question_kind", "mcq"),
                "cognitive_action": q.get("cognitive_action", ""),
                "timestamp": format_seconds(q.get("timestamp_seconds", "")),
                "timestamp_seconds": q.get("timestamp_seconds", ""),
                "reveal_after": format_seconds(q.get("reveal_after_seconds", "")),
                "reveal_after_seconds": q.get("reveal_after_seconds", ""),
                "stem": q.get("stem", ""),
                "option_A": opt_lookup.get("A", ""),
                "option_B": opt_lookup.get("B", ""),
                "option_C": opt_lookup.get("C", ""),
                "option_D": opt_lookup.get("D", ""),
                "correct_option": q.get("correct_option", ""),
                "expected_answer": q.get("expected_answer", ""),
                "marking_points": " | ".join(q.get("marking_points", []) or []),
                "feedback_correct": q.get("feedback_correct", ""),
                "feedback_incorrect": q.get("feedback_incorrect", ""),
                "rationale": q.get("rationale", ""),
                "learning_objective": q.get("learning_objective", ""),
                "difficulty": q.get("difficulty", ""),
                "evidence_start": format_seconds(q.get("evidence_start_seconds", "")),
                "evidence_start_seconds": q.get("evidence_start_seconds", ""),
                "evidence_end": format_seconds(q.get("evidence_end_seconds", "")),
                "evidence_end_seconds": q.get("evidence_end_seconds", ""),
                "placement_reason": q.get("placement_reason", ""),
                "single_best_answer": qf.get("single_best_answer", ""),
                "distractors_plausible": qf.get("distractors_plausible", ""),
                "conceptual_not_trivia": qf.get("conceptual_not_trivia", ""),
                "aligned_to_video": qf.get("aligned_to_video", ""),
                "teacher_review_recommended": qf.get("teacher_review_recommended", ""),
            }
        )


    pre_questions = video_mcq_result.get("pre_questions", []) or []

    if not pre_questions and video_mcq_result.get("pre_question"):
        pre_questions = [video_mcq_result.get("pre_question")]

    for q in pre_questions:
        add_question(q)

    for q in video_mcq_result.get("embedded_questions", []) or []:
        add_question(q)

    return rows

#Streamlit frontend
st.set_page_config(layout="wide", page_title="Vox-LM SAQ Marking Prototype for Vox 2.0")
st.title(":blue[Vox-LM] _Prototype_")
tab_marking, tab_summary, tab_student_reports, tab_mcq_from_videos = st.tabs(
    [
        "SAQ Marking", "Class Performance Summary", "Student Test Reports", "Video Question Generator"]
)


#defaults for frontend state
if "grade_result" not in st.session_state:
    st.session_state.grade_result = None

if "override_mode" not in st.session_state:
    st.session_state.override_mode = False

if "q_edit_mode" not in st.session_state:
    st.session_state.q_edit_mode = True

if "discipline_choice" not in st.session_state:
    st.session_state.discipline_choice = "Dentistry"

discipline = st.session_state.discipline_choice.lower()

if "batch_criterion_results_df" not in st.session_state:
    st.session_state.batch_criterion_results_df = None

#if "batch_norm_results_df" not in st.session_state:
 #   st.session_state.batch_norm_results_df = None

if "batch_norm_teacher_results_df" not in st.session_state:
    st.session_state.batch_norm_teacher_results_df = None

if "batch_norm_diagnostic_results_df" not in st.session_state:
    st.session_state.batch_norm_diagnostic_results_df = None


#defaults for challenge mode state
if "challenge_mode" not in st.session_state:
    st.session_state.challenge_mode = False

if "challenge_used" not in st.session_state:
    st.session_state.challenge_used = False

if "challenge_submitted" not in st.session_state:
    st.session_state.challenge_submitted = False

if "challenge_reason_text" not in st.session_state:
    st.session_state.challenge_reason_text = ""

if "last_grade_payload" not in st.session_state:
    st.session_state.last_grade_payload = None

if "original_grade_result" not in st.session_state:
    st.session_state.original_grade_result = None

if "last_has_subquestions" not in st.session_state:
    st.session_state.last_has_subquestions = False

#summary session state defaults
if "batch_summary_result" not in st.session_state:
    st.session_state.batch_summary_result = None

if "batch_summary_debug_prompt" not in st.session_state:
    st.session_state.batch_summary_debug_prompt = ""

#student report session state defaults
if "student_reports_result" not in st.session_state:
    st.session_state.student_reports_result = None

if "student_reports_debug_prompt" not in st.session_state:
    st.session_state.student_reports_debug_prompt = ""

if "last_handwritten_upload_sig" not in st.session_state:
    st.session_state.last_handwritten_upload_sig = None

if "handwriting_transcription_confidence" not in st.session_state:
    st.session_state.handwriting_transcription_confidence = None

if "handwriting_transcription_debug" not in st.session_state:
    st.session_state.handwriting_transcription_debug = ""

if "teacher_finalised" not in st.session_state:
    st.session_state.teacher_finalised = False

if "finalised_grade_result" not in st.session_state:
    st.session_state.finalised_grade_result = None

#chat with Vox-LM session state defaults

if "voxlm_chat_open" not in st.session_state:
    st.session_state.voxlm_chat_open = False

if "voxlm_chat_history" not in st.session_state:
    st.session_state.voxlm_chat_history = []

if "voxlm_chat_input" not in st.session_state:
    st.session_state.voxlm_chat_input = ""

#refine model answer session state defaults

if "refine_model_answer_open" not in st.session_state:
    st.session_state.refine_model_answer_open = False

if "refine_model_answer_result" not in st.session_state:
    st.session_state.refine_model_answer_result = None

if "refine_model_answer_debug_prompt" not in st.session_state:
    st.session_state.refine_model_answer_debug_prompt = ""

if "refine_model_answer_text" not in st.session_state:
    st.session_state.refine_model_answer_text = ""

if "pending_refined_model_answer" not in st.session_state:
    st.session_state.pending_refined_model_answer = None

if "video_mcq_result" not in st.session_state:
    st.session_state.video_mcq_result = None

if "video_mcq_debug_prompt" not in st.session_state:
    st.session_state.video_mcq_debug_prompt = ""

if "video_mcq_session_counter" not in st.session_state:
    st.session_state.video_mcq_session_counter = 0

if "video_mcq_display_id" not in st.session_state:
    st.session_state.video_mcq_display_id = ""

#sidebar
with st.sidebar:
    st.header("Vox-LM Controls")

    discipline_label= st.radio(
        "Faculty",
        ["Dentistry", "Medicine", "Law", "Education"],
        index=["Dentistry", "Medicine", "Law", "Education"].index(st.session_state.discipline_choice),
        key="discipline_choice",
    )

    if st.button("**:red[Teacher Override]**", key="btn_override"):
        st.session_state.override_mode = not st.session_state.override_mode

    st.write(f"Override mode: {'ON' if st.session_state.override_mode else 'OFF'}")

    if st.button("**:blue[Question Editor]**", key="btn_qedit"):
        st.session_state.q_edit_mode = not st.session_state.q_edit_mode

    st.write(f"Edit question/rubric mode: {'ON' if st.session_state.q_edit_mode else 'OFF'}")
    st.markdown("---")
    st.header("Batch grading")

    batch_csv = st.file_uploader(
        "**:green[Upload CSV of student responses. CSV should have two columns and start with 'STUDENT ID' (first column) and 'STUDENT RESPONSES' (second column)]**",
        type=["csv"],
        key="batch_csv_uploader"
    )

    batch_grade_btn = st.button("**Grade uploaded CSV**", key="btn_batch_grade")

discipline = st.session_state.discipline_choice.lower()
override = st.session_state.override_mode
q_edit_mode = st.session_state.q_edit_mode

with tab_marking:
#three column panel
    col_left, col_mid, col_right = st.columns(3)

    #left
    with col_left:
        st.subheader(":violet[Question, Rubric, Model Answer, and Examples]")

        has_subquestions = st.checkbox("**:blue[Question has subquestions]**", value=False)

        disabled_left = not st.session_state.q_edit_mode

        question_stem = st.text_area(
            ":blue[Question stem]",
            height=150,
            disabled=disabled_left,
        )

        if st.session_state.pending_refined_model_answer is not None:
            st.session_state.model_answer_input = st.session_state.pending_refined_model_answer
            st.session_state.pending_refined_model_answer = None

        model_answer = st.text_area(
            ":blue[Model answer]",
            height=100,
            disabled=disabled_left,
            key="model_answer_input"
        )

        global_rubric = st.text_area(
            ":blue[Rubric / marking scheme (optional)]",
            height=100,
            disabled=disabled_left,
        )

        max_score = st.number_input(
            ":blue[Total score]",
            min_value=1.0,
            max_value=100.0,
            value=10.0,
            step=0.5,
            disabled=disabled_left,
        )

        if has_subquestions:
            st.markdown("**:blue[Subquestions format]**")
            st.caption(
                "Format each subquestion as:  \n"
                "Number and colon (e.g., 1:, 2:, 3:, etc.)  \n"
                "Question text  \n"
                "Sub-score for question  \n"
                "Rubric/model answer\n\n"
                "Start separate subquestions with a blank line.\n"
                "First line may also be '1', '1.' or '1)'."
            )

            subq_text = st.text_area(
                ":blue[Subquestions]",
                height=220,
                disabled=disabled_left,
            )
        else:
            subq_text = ""
            st.info("Subquestions disabled; sub-scores will show as NA.")
        
        parsed_subquestions: List[Dict] = []
        if has_subquestions and subq_text.strip():
            try:
                parsed_subquestions = parse_subquestions_from_text(subq_text)
                if not parsed_subquestions:
                    st.warning("No valid subquestions parsed. Check the text format.")
            except Exception as e:
                st.error(f"Error parsing subquestions text: {e}")
                parsed_subquestions = []

        if st.button("**:blue[Refine model answers]**", key="btn_refine_model_answer"):
            st.session_state.refine_model_answer_open = not st.session_state.refine_model_answer_open

#MIGHT REMOVE BASED ON UI PERFORMANCE
            if st.session_state.refine_model_answer_open:
                st.session_state.refine_model_answer_text = model_answer or ""
                st.session_state.refine_model_answer_result = None
                st.session_state.refine_model_answer_debug_prompt = ""

        if st.session_state.refine_model_answer_open:
            st.markdown("### Refine model answer for Vox-LM marking")
            st.caption(
                "Paste or edit the model answer below. Vox-LM will rate if it is suitable for marking student responses to this question"
                "and suggest rubrics as well. Please note that this does not automatically replace your model answer."
            )

            if not st.session_state.refine_model_answer_text:
                st.session_state.refine_model_answer_text = model_answer

            st.text_area(
                "Model answer to review",
                key="refine_model_answer_text",
                height=220,
            )

            if st.button(":blue[Refine model answer]", key="btn_submit_refine_model_answer"):
                try:
                    refine_question = {
                        "exam_id": "EXAM",
                        "question_id": "Q1",
                        "stem": question_stem,
                        "max_score": max_score,
                        "subquestions": parsed_subquestions if has_subquestions else [],
                        "model_answer": st.session_state.refine_model_answer_text,
                        "rubric": global_rubric,
                    }

                    refine_payload = {
                        "question": refine_question,
                        "discipline": discipline,
                    }

                    with st.spinner("Vox-LM is reviewing the model answer..."):
                        res = requests.post(
                            BACKEND_REFINE_MODEL_ANSWER_URL,
                            json=refine_payload,
                            headers={"x-api-key": BACKEND_API_KEY},
                            timeout=300,
                        )

                    if res.status_code != 200:
                        st.error(f"Model answer refinement backend error: {res.status_code} {res.text}")
                    else:
                        data = res.json()
                        st.session_state.refine_model_answer_result = data
                        st.session_state.refine_model_answer_debug_prompt = data.get("debug_prompt", "")
                        st.success("Model answer review completed.")

                except Exception as e:
                    st.error(f"Failed to refine model answer: {e}")

            refine_result = st.session_state.refine_model_answer_result

            if refine_result:
                st.markdown("#### Model answer rating")
                st.metric(
                    "Rating",
                    f"{float(refine_result.get('rating_score', 0.0)):.1f}/100",
                    refine_result.get("rating_label", ""),
                )

                st.markdown("#### Strengths")
                strengths = refine_result.get("strengths", []) or []
                if strengths:
                    for item in strengths:
                        st.write(f"- {item}")
                else:
                    st.write("No major strengths identified.")

                st.markdown("#### Issues")
                issues = refine_result.get("issues", []) or []
                if issues:
                    for item in issues:
                        st.write(f"- {item}")
                else:
                    st.write("No major issues identified.")

                st.markdown("#### Rewritten model answer")
                st.text_area(
                    "Copy this into the main model answer box if suitable",
                    value=refine_result.get("rewritten_model_answer", ""),
                    height=220,
                    key="refined_model_answer_output",
                )

                st.markdown("#### Suggested rubrics")
                st.text_area(
                    "Suggested marking structure",
                    value=refine_result.get("suggested_marking_structure", ""),
                    height=220,
                    key="refined_marking_structure_output",
                )

                if st.button("Use refined model answer", key="btn_use_refined_model_answer"):
                    st.session_state.pending_refined_model_answer = refine_result.get("rewritten_model_answer", "")
                    st.success("Rewritten model answer copied into the main model answer box.")
                    st.rerun()
    

        images_files = st.file_uploader(
            ":blue[Upload SAQ images (optional)]",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            disabled=disabled_left,
        )

        st.markdown("**:blue[Few-shot examples (optional)]**")
        st.caption(
            "Each example:  \n"
            "Student Number: [ID]  \n"
            "Answer: [student answers]  \n"
            "Marker score: [number]\n\n"
            "**Separate examples with a blank line.**"
        )

        few_shot_text = st.text_area(
            ":blue[Enter examples here]",
            height=220,
            disabled=disabled_left,
        )

    #middle panel
    with col_mid:
        st.subheader(":violet[Student Response]")

        handwritten_response_mode = st.checkbox(
            "**:blue[Handwritten response]**",
            key="handwritten_response_mode",
            disabled=st.session_state.challenge_used,
        )

        if handwritten_response_mode:
            handwritten_response_file = st.file_uploader(
                "Upload handwritten image of student response",
                type=["png", "jpg", "jpeg"],
                key="handwritten_response_image",
                disabled=st.session_state.challenge_used,
            )

            if handwritten_response_file is not None:
                file_bytes = handwritten_response_file.getvalue()
                file_sig = hashlib.md5(file_bytes).hexdigest()

                if st.session_state.last_handwritten_upload_sig != file_sig:
                    try:
                        with st.spinner("Transcribing handwritten response..."):
                            payload = {
                                "image": base64.b64encode(file_bytes).decode("utf-8"),
                                "discipline": discipline,
                            }

                            res = requests.post(
                                BACKEND_TRANSCRIBE_URL,
                                json=payload,
                                headers={"x-api-key": BACKEND_API_KEY},
                                timeout=300,
                            )

                            if res.status_code != 200:
                                st.error(f"Transcription backend error: {res.status_code} {res.text}")
                            else:
                                data = res.json()
                                st.session_state.student_answer = data.get("transcription", "")
                                st.session_state.handwriting_transcription_confidence = data.get("confidence", 0.0)
                                st.session_state.handwriting_transcription_debug = data.get("debug_prompt", "")
                                st.session_state.last_handwritten_upload_sig = file_sig
                                st.success("Handwritten response transcribed. Please inspect before grading and edit if necessary.")
                    except Exception as e:
                        st.error(f"Failed to transcribe handwritten response: {e}")

            if st.session_state.handwriting_transcription_confidence is not None:
                st.caption(
                    f"HTR confidence: {float(st.session_state.handwriting_transcription_confidence):.1f} %"
                )

        student_answer_text = st.text_area(
            "Answer box",
            height=250,
            key="student_answer",
            disabled=st.session_state.challenge_used,
        )

        response_id = st.text_input(
            "Student name (optional)",
            value="John Doe",
            disabled=st.session_state.challenge_used,
        )


        if st.button("**:blue[Grade student]**"):
            # Build question object
            question_dict = {
                "exam_id": "EXAM",
                "question_id": "Q1",
                "stem": question_stem,
                "max_score": max_score,
                "subquestions": parsed_subquestions if has_subquestions else [],
                "model_answer": model_answer,
                "rubric": global_rubric,
            }

            # Few-shot list from text
            few_shot_list: List[Dict] = []
            if few_shot_text.strip():
                try:
                    few_shot_list = parse_few_shot_from_text(
                        few_shot_text,
                        has_subquestions=has_subquestions,
                        subquestions=parsed_subquestions,
                    )
                except Exception as e:
                    st.error(f"Error parsing few-shot examples: {e}")
                    few_shot_list = []

            # Images in base64 form
            images_b64: List[str] = []
            if images_files:
                for f in images_files:
                    b = f.getvalue()
                    images_b64.append(base64.b64encode(b).decode("utf-8"))

            # Build payload and call backend
            if has_subquestions and parsed_subquestions:
                answers = parse_student_answers_from_text(
                    student_answer_text, parsed_subquestions
                )
            else:
                answers = {"overall": student_answer_text}

            payload = {
                "question": question_dict,
                "few_shot": few_shot_list,
                "student_response": {
                    "response_id": response_id,
                    "answers": answers,
                },
                "has_subquestions": has_subquestions,
                "images": images_b64,
                "discipline": discipline,
                "challenge_mode": False,
                "challenge_reason": None,
                "original_total_score": None,
                "original_sub_scores": None,
            }

            # Reset challenge state for a fresh grading run
            st.session_state.challenge_mode = False
            st.session_state.challenge_used = False
            st.session_state.challenge_submitted = False
            st.session_state.challenge_reason_text = ""
            st.session_state.last_grade_payload = copy.deepcopy(payload)
            st.session_state.original_grade_result = None
            st.session_state.last_has_subquestions = has_subquestions
            st.session_state.teacher_finalised = False
            st.session_state.finalised_grade_result = None

            st.session_state.voxlm_chat_open = False
            st.session_state.voxlm_chat_history = []
            st.session_state.voxlm_chat_input = ""


            try:
                res = requests.post(
                    BACKEND_GRADE_URL,
                    json=payload,
                    headers={"x-api-key": BACKEND_API_KEY},
                    timeout=300
                )
                if res.status_code != 200:
                    st.error(f"Backend error: {res.status_code} {res.text}")
                else:
                    data = res.json()
                    st.session_state.grade_result = data
                    st.session_state.original_grade_result = copy.deepcopy(data)
                    st.session_state.debug_prompt_view = data.get("debug_prompt", "")
            except Exception as e:
                st.error(f"Request to backend failed: {e}")

        # Show highlight + challenge/reset only AFTER grading
        if st.session_state.grade_result is not None:
            st.markdown("---")
            st.markdown("**Highlighted student answer**")

            result = st.session_state.grade_result
            highlights = result.get("highlights", {})

            all_correct: List[str] = []
            all_out_of_scope: List[str] = []
            all_misconception: List[str] = []
            all_uncertain: List[str] = []

            for sid, h in (highlights or {}).items():
                all_correct.extend(h.get("correct", []) or [])
                all_out_of_scope.extend(h.get("out_of_scope", []) or [])

                # Backward compatibility with old backend outputs.
                all_misconception.extend(h.get("misconception", []) or h.get("incorrect", []) or [])

                all_uncertain.extend(h.get("uncertain", []) or [])

            st.caption(
                "**Highlight key:** "
                "🟢 = Correct; "
                "🟡 = Out of scope; "
                "🟠 = Misconception; "
                "🔵 = Uncertain"
            )

            html = build_highlighted_html(
                student_answer_text,
                all_correct,
                all_out_of_scope,
                all_misconception,
                all_uncertain,
            )


            st.markdown(html, unsafe_allow_html=True)

            st.markdown("---")

            # Challenge button shown only before challenge is used
            if st.session_state.teacher_finalised:
                st.info("Grade is finalised. Re-open it to use challenge or edit functions.")
            elif not st.session_state.challenge_used:

                st.button(
                    "**:red[CHALLENGE!]**",
                    key="btn_challenge_score",
                    on_click=start_challenge,
                )

            # Challenge in progress: show textbox + submit once
            elif st.session_state.challenge_mode and not st.session_state.challenge_submitted:
                st.info(
                    "Challenge mode has locked the original submitted answers. It can no longer be edited."
                )

                st.markdown("**Challenge score**")
                st.caption(
                    "Enter reasons why you believe you deserve a higher score. "
                    "The model will only award more marks if your challenge is supported by your original submitted answer."
                )

                st.text_area(
                    "Challenge Vox-LM's reasons and grading",
                    key="challenge_reason_text",
                    height=180,
                )

                if st.button("Submit challenge", key="btn_submit_challenge"):
                    if st.session_state.challenge_submitted:
                        st.warning("Challenge has already been submitted.")
                    elif not st.session_state.challenge_reason_text.strip():
                        st.warning("Please enter your challenge reasons before submitting.")
                    elif (
                        st.session_state.last_grade_payload is None
                        or st.session_state.original_grade_result is None
                    ):
                        st.error("Original grade context not found. Please grade the student first.")
                    else:
                        challenge_payload = copy.deepcopy(st.session_state.last_grade_payload)
                        challenge_payload["challenge_mode"] = True
                        challenge_payload["challenge_reason"] = st.session_state.challenge_reason_text
                        challenge_payload["original_total_score"] = (
                            st.session_state.original_grade_result.get("total_score", 0.0)
                        )
                        challenge_payload["original_sub_scores"] = (
                            st.session_state.original_grade_result.get("sub_scores", {})
                        )

                        try:
                            res = requests.post(
                                BACKEND_GRADE_URL,
                                json=challenge_payload,
                                headers={"x-api-key": BACKEND_API_KEY},
                                timeout=300
                            )
                            if res.status_code != 200:
                                st.error(f"Challenge request failed: {res.status_code} {res.text}")
                            else:
                                data = res.json()
                                st.session_state.grade_result = data
                                st.session_state.challenge_submitted = True
                                st.session_state.challenge_mode = False
                                st.session_state.debug_prompt_view = data.get("debug_prompt", "")
                                st.success("Challenge submitted. The answer has been re-marked.")
                        except Exception as e:
                            st.error(f"Challenge request failed: {e}")

            elif st.session_state.challenge_submitted:
                st.error("Challenge already submitted. Only one challenge is allowed!")

            st.button(
                "**:green[Reset]**",
                key="btn_reset_student",
                on_click=reset_student_state,
            )

    if st.session_state.grade_result is not None:
        st.markdown("---")

        if st.button("**:blue[Chat with Vox-LM]**", key="btn_open_voxlm_chat"):
            st.session_state.voxlm_chat_open = not st.session_state.voxlm_chat_open

        if st.session_state.voxlm_chat_open:
            st.markdown("###:violet[Chat with Vox-LM]")
            st.caption(
                "This chat helps explain feedback and guide learning. "
                "It does not change the grade."
            )

            for msg in st.session_state.voxlm_chat_history:
                if msg.get("role") == "user":
                    st.chat_message("user").write(msg.get("content", ""))
                else:
                    st.chat_message("assistant", avatar="🧠").write(msg.get("content", ""))

            chat_input = st.text_input(
                "Ask Vox-LM about your feedback and next steps for improvement",
                key="voxlm_chat_input",
                placeholder="Example question: Why did I lose marks?",
            )

            if st.button(":blue[Send to Vox-LM]", key="btn_send_voxlm_chat"):
                if not chat_input.strip():
                    st.warning("Please enter a question.")
                else:
                    try:
                        user_msg = chat_input.strip()
                        previous_history = list(st.session_state.voxlm_chat_history)

                        chat_payload = {
                            "question": {
                                "exam_id": "EXAM",
                                "question_id": "Q1",
                                "stem": question_stem,
                                "max_score": max_score,
                                "subquestions": parsed_subquestions if has_subquestions else [],
                                "model_answer": model_answer,
                                "rubric": global_rubric,
                            },
                            "student_response": {
                                "response_id": response_id,
                                "answers": (
                                    parse_student_answers_from_text(student_answer_text, parsed_subquestions)
                                    if has_subquestions and parsed_subquestions
                                    else {"overall": student_answer_text}
                                ),
                            },
                            "grade_result": st.session_state.grade_result,
                            "chat_history": previous_history,
                            "user_message": user_msg,
                            "discipline": discipline,
                        }

                        with st.spinner("Vox-LM is responding..."):
                            res = requests.post(
                                BACKEND_CHAT_URL,
                                json=chat_payload,
                                headers={"x-api-key": BACKEND_API_KEY},
                                timeout=300,
                            )

                        if res.status_code != 200:
                            st.error(f"Chat backend error: {res.status_code} {res.text}")
                        else:
                            data = res.json()
                            assistant_message = data.get("assistant_message", "")

                            st.session_state.voxlm_chat_history.append({
                                "role": "user",
                                "content": user_msg,
                            })

                            st.session_state.voxlm_chat_history.append({
                                "role": "assistant",
                                "content": assistant_message,
                            })

                            st.rerun()

                    except Exception as e:
                        st.error(f"Chat request failed: {e}")

            if st.button(":blue[Clear chat]", key="btn_clear_voxlm_chat"):
                st.session_state.voxlm_chat_history = []
                st.rerun()



    #bATCH GRADING FRONT

    if batch_csv is not None and batch_grade_btn:
        st.session_state.batch_criterion_results_df = None
        st.session_state.batch_norm_teacher_results_df = None
        st.session_state.batch_norm_diagnostic_results_df = None
        try:
            try:
                df = pd.read_csv(batch_csv, dtype_backend="numpy_nullable")
            except TypeError:
                df = pd.read_csv(batch_csv)

            response_cols = [c for c in df.columns if str(c).startswith("STUDENT RESPONSES")]
            if not response_cols:
                st.error("CSV must contain a column whose title starts with 'STUDENT RESPONSES'.")
            else:
                response_col = response_cols[0]
                has_student_id_col = "STUDENT_ID" in df.columns

                question_dict = {
                    "exam_id": "EXAM",
                    "question_id": "Q1",
                    "stem": question_stem,
                    "max_score": max_score,
                    "subquestions": parsed_subquestions if has_subquestions else [],
                    "model_answer": model_answer,
                    "rubric": global_rubric,
                }

                few_shot_list: List[Dict] = []
                if few_shot_text.strip():
                    try:
                        few_shot_list = parse_few_shot_from_text(
                            few_shot_text,
                            has_subquestions=has_subquestions,
                            subquestions=parsed_subquestions,
                        )
                    except Exception as e:
                        st.error(f"Error parsing few-shot examples: {e}")
                        few_shot_list = []

                images_b64: List[str] = []
                if images_files:
                    for f in images_files:
                        b = f.getvalue()
                        images_b64.append(base64.b64encode(b).decode("utf-8"))

                results_df = df.copy()
                results_df["total_score"] = None
                results_df["confidence"] = None
                results_df["rationale"] = None
                results_df["feedback_overall"] = None
                results_df["needs_review"] = None
                results_df["review_reasons"] = None
                results_df["missing_key_point"] = None

                # Pre-create sub-score columns if subquestions exist
                if has_subquestions and parsed_subquestions:
                    for sq in parsed_subquestions:
                        sid = str(sq.get("id"))
                        col_name = f"sub_{sid}"
                        results_df[col_name] = None

                total_rows = len(df)
                progress = st.progress(0.0)
                status_text = st.empty()

                for i, row in df.iterrows():
                    raw_response = row[response_col]
                    student_text = "" if pd.isna(raw_response) else str(raw_response)

                    # Student ID
                    if has_student_id_col:
                        raw_id = row["STUDENT_ID"]
                        student_id = str(raw_id) if pd.notna(raw_id) else f"student_{i+1}"
                    else:
                        student_id = f"student_{i+1}"

                    # Build answers dict for this student
                    if has_subquestions and parsed_subquestions:
                        answers = parse_student_answers_from_text(
                            student_text, parsed_subquestions
                        )
                    else:
                        answers = {"overall": student_text}

                    payload = {
                        "question": question_dict,
                        "few_shot": few_shot_list,
                        "student_response": {
                            "response_id": student_id,
                            "answers": answers,
                        },
                        "has_subquestions": has_subquestions,
                        "images": images_b64,
                        "discipline": discipline,
                        "challenge_mode": False,
                        "challenge_reason": None,
                        "original_total_score": None,
                        "original_sub_scores": None,
                    }

    #Call backend
                    try:
                        res = requests.post(BACKEND_GRADE_URL, json=payload, headers={"x-api-key": BACKEND_API_KEY}, timeout=300)
                        if res.status_code != 200:
                            st.warning(f"Row {i}: backend error {res.status_code}")
                            result_json = None
                        else:
                            result_json = res.json()
                    except Exception as e:
                        st.warning(f"Row {i}: request failed: {e}")
                        result_json = None

    #Fill result columns
                    if result_json is not None:
                        results_df.at[i, "total_score"] = result_json.get("total_score")
                        results_df.at[i, "confidence"] = result_json.get("confidence")
                        results_df.at[i, "rationale"] = result_json.get("rationale")

                        fb = result_json.get("feedback", {}) or {}
                        results_df.at[i, "feedback_overall"] = fb.get("_overall", "")
                        results_df.at[i, "needs_review"] = result_json.get("needs_review", False)
                        results_df.at[i, "review_reasons"] = " | ".join(result_json.get("review_reasons", []) or [])
                        results_df.at[i, "missing_key_point"] = result_json.get(
                            "missing_key_point", ""
                        )

                        results_df.at[i, "feedback_json"] = json.dumps(fb, ensure_ascii=False)

                        for k, v in fb.items():
                            if k == "_overall":
                                continue
                            col_name = f"feedback_{k}"
                            if col_name not in results_df.columns:
                                results_df[col_name] = None
                            results_df.at[i, col_name] = v

                        results_df.at[i, "highlights_json"] = json.dumps(
                            result_json.get("highlights", {}) or {},
                            ensure_ascii=False
                        )

                        #fb = result_json.get("feedback", {}) or {}
                        #results_df.at[i, "feedback_overall"] = fb.get("_overall", "")

                        sub_scores = result_json.get("sub_scores", {}) or {}
                        for sid, s_val in sub_scores.items():
                            col_name = f"sub_{sid}"
                            if col_name not in results_df.columns:
                                results_df[col_name] = None
                            results_df.at[i, col_name] = s_val

                    progress.progress((i + 1) / max(1, total_rows))
                    status_text.text(f"Graded {i+1}/{total_rows} student responses")

#Store in session state                
                results_df["result_type"] = "rubrics_model_answer_referenced"
                st.session_state.batch_criterion_results_df = sanitize_df_for_streamlit(results_df)

                criterion_csv_text = results_df.to_csv(index=False)

                norm_payload = {
                    "csv_text": criterion_csv_text,
                    "question": question_dict,
                    "has_subquestions": has_subquestions,
                    "discipline": discipline,
                }

                norm_res = requests.post(
                    BACKEND_NORM_URL,
                    json=norm_payload,
                    headers={"x-api-key": BACKEND_API_KEY},
                    timeout=300
                )

                if norm_res.status_code != 200:
                    st.warning(f"Norm-referenced analysis failed: {norm_res.status_code} {norm_res.text}")
                    st.session_state.batch_norm_teacher_results_df = None
                    st.session_state.batch_norm_diagnostic_results_df = None
                else:
                    norm_data = norm_res.json()

                    teacher_df = pd.DataFrame(norm_data.get("teacher_rows", []))
                    diagnostic_df = pd.DataFrame(norm_data.get("diagnostic_rows", []))

                    st.session_state.batch_norm_teacher_results_df = sanitize_df_for_streamlit(teacher_df)
                    st.session_state.batch_norm_diagnostic_results_df = sanitize_df_for_streamlit(diagnostic_df)

                st.success("Batch grading completed.")

        except Exception as e:
            st.error(f"Failed to process batch CSV: {e}")


    #Right panel
    with col_right:
        st.subheader(":violet[Results]")

        result = st.session_state.grade_result

        if result is None:
            st.info("Run grading to see student results.")
        else:
            override = st.session_state.override_mode
            effective_override = override and not st.session_state.teacher_finalised

            total_score = result.get("total_score", 0.0)
            sub_scores = result.get("sub_scores", {})
            rationale = result.get("rationale", "")
            feedback = result.get("feedback", {})
            confidence = result.get("confidence", 0.0)
            challenge_review = result.get("challenge_review", "")
            original_total_score_before_challenge = result.get("original_total_score", None)

#review status front
            st.markdown("#### Review status")

            combined_review_reasons = list(result.get("review_reasons", []) or [])

            if st.session_state.handwriting_transcription_confidence is not None:
                try:
                    if float(st.session_state.handwriting_transcription_confidence) < 70:
                        combined_review_reasons.append("Low handwriting transcription confidence")
                except Exception:
                    pass

            combined_review_reasons = list(dict.fromkeys(combined_review_reasons))
            needs_review_display = bool(result.get("needs_review", False)) or bool(combined_review_reasons)

            if needs_review_display:
                st.warning("Needs Review!")
                for reason in combined_review_reasons:
                    st.write(f"- {reason}")
            else:
                st.success("No major issues identified with the VOX-LM grading.")

#mark approval front
            st.markdown("#### Mark approval status")

            if st.session_state.teacher_finalised:
                st.success("Status: Approved by teacher")
                st.caption("Editing is locked while the grade is being finalised.")
                if st.button("Reopen grade", key="btn_reopen_mark"):
                    st.session_state.teacher_finalised = False
                    st.session_state.finalised_grade_result = None
                    st.rerun()
            else:
                st.info("Status: Vox-LM provisional grade (not yet finalised by teacher)")
                if st.button("**:blue[Approve grade]**", key="btn_finalise_mark"):
                    st.session_state.teacher_finalised = True
                    st.session_state.finalised_grade_result = copy.deepcopy(result)
                    st.rerun()

#Total score front
            if effective_override:
                total_score = st.number_input(
                    "Total score",
                    value=float(total_score),
                    step=0.5,
                    key="override_total_score",
                )
            else:
                st.metric("Total score", f"{total_score}")

            #subscore front
            st.markdown("### Sub-scores")
            if not st.session_state.last_has_subquestions:
                st.write("Sub-scores: NA (no subquestions enabled)")
            else:
                if isinstance(sub_scores, dict):
                    for sid in sub_scores.keys():
                        val = sub_scores[sid]
                        if effective_override:
                            try:
                                default_val = float(val)
                            except Exception:
                                default_val = 0.0
                            sub_scores[sid] = st.number_input(
                                f"Sub-score ({sid})",
                                value=default_val,
                                step=0.5,
                                key=f"override_subscore_{sid}",
                            )
                        else:
                            st.write(f"({sid}): {val}")
                else:
                    st.write("Invalid sub_scores format from backend")

            #Rationale front
            st.markdown("### Rationale / Reasoning")
            if effective_override:
                rationale = st.text_area(
                    "Rationale", value=rationale, height=100, key="override_rationale"
                )
            else:
                st.write(rationale)
#Missing key point front
            st.markdown("### Missing key point")
            missing_key_point = result.get("missing_key_point", "")

            if effective_override:
                missing_key_point = st.text_area(
                    "Missing key point",
                    value=missing_key_point,
                    height=100,
                    key="override_missing_key_point",
                )
            else:
                st.write(missing_key_point or "No major missing point identified.")


            #Feedback and improvement front
            st.markdown("### Feedback & Areas for improvement")
            if effective_override:
                feedback_text = st.text_area(
                    "Feedback (JSON)",
                    value=json.dumps(feedback, indent=2),
                    height=150,
                    key="override_feedback",
                )
                try:
                    feedback = json.loads(feedback_text)
                except Exception:
                    st.error("Invalid feedback JSON; keeping original.")
            else:
                overall_fb = feedback.get("_overall")
                if overall_fb:
                    st.write(f"Overall: {overall_fb}")
                for k, v in feedback.items():
                    if k == "_overall":
                        continue
                    st.write(f"({k}) {v}")

            #Confidence slide bar
            st.markdown("#### Grading Confidence")
            if effective_override:
                confidence = st.slider(
                    "Confidence",
                    min_value=0.0,
                    max_value=100.0,
                    value=float(confidence),
                    step=0.05,
                    key="override_confidence",
                )
            else:
                st.progress(min(1.0, max(0.0, confidence / 100.0)))
                st.write(f"{confidence:.1f}% confidence in this grading")

# challenge review front 
            if challenge_review:
                st.markdown("#### Challenge Review")
                if original_total_score_before_challenge is not None:
                    st.write(f"Original score before challenge: {original_total_score_before_challenge}")
                    try:
                        delta = float(total_score) - float(original_total_score_before_challenge)
                        if delta > 0:
                            st.write(f"Score increase after challenge: +{delta:.1f}")
                        else:
                            st.write("Score unchanged after challenge.")
                    except Exception:
                        pass
                st.write(challenge_review)

            #Human overide information items are saved after teacher edits
            if effective_override:
                result["total_score"] = total_score
                result["sub_scores"] = sub_scores
                result["rationale"] = rationale
                result["feedback"] = feedback
                result["confidence"] = confidence
                result["missing_key_point"] = missing_key_point
                st.session_state.grade_result = result

#model answer-referenced batch results
        st.markdown("---")
        st.subheader(":violet[Model answer-referenced Batch results]")

        if st.session_state.batch_criterion_results_df is not None:
            criterion_df = sanitize_df_for_streamlit(st.session_state.batch_criterion_results_df)
            st.dataframe(criterion_df, use_container_width=True)

            criterion_csv_buffer = io.StringIO()
            criterion_df.to_csv(criterion_csv_buffer, index=False)

            st.download_button(
                ":green[Download rubrics/model answer graded CSV]",
                data=criterion_csv_buffer.getvalue(),
                file_name="voxlm_batch_model_answer_results.csv",
                mime="text/csv",
            )
        else:
            st.info("Upload a CSV and click 'Grade uploaded CSV' in the SIDE BAR to see batch results based on rubrics/model answers here.")

        st.subheader(":violet[Norm-referenced Batch results]")

        st.caption(
            "This norm-referenced output compares each response with the rest of the class. "
            "It does not change the rubrics/model answer score."
        )


        if st.session_state.batch_norm_teacher_results_df is not None:
            norm_teacher_df = sanitize_df_for_streamlit(st.session_state.batch_norm_teacher_results_df)
            st.dataframe(norm_teacher_df, use_container_width=True)

            norm_teacher_csv_buffer = io.StringIO()
            norm_teacher_df.to_csv(norm_teacher_csv_buffer, index=False)

            st.download_button(
                ":green[Download norm-referenced CSV]",
                data=norm_teacher_csv_buffer.getvalue(),
                file_name="voxlm_batch_norm_referenced.csv",
                mime="text/csv",
            )

            with st.expander("OPTIONAL: Show advanced norm-referenced diagnostics"):
                st.caption(
                    "These technical metrics are optional diagnostics used to support the qualitative comparisons above. "
                    "They are not intended to replace teacher judgement."
                )
                if st.session_state.batch_norm_diagnostic_results_df is not None:
                    norm_diag_df = sanitize_df_for_streamlit(st.session_state.batch_norm_diagnostic_results_df)
                    st.dataframe(norm_diag_df, use_container_width=True)

                    norm_diag_csv_buffer = io.StringIO()
                    norm_diag_df.to_csv(norm_diag_csv_buffer, index=False)

                    st.download_button(
                        "Download advanced diagnostics CSV",
                        data=norm_diag_csv_buffer.getvalue(),
                        file_name="voxlm_batch_norm_referenced_diagnostics.csv",
                        mime="text/csv",
                    )
                else:
                    st.info("No diagnostic norm-referenced data available.")
        else:
            st.info("Norm-referenced batch output will appear after rubrics/model answer batch grading completes.")



    #sidebar 2
    with st.sidebar:
        if "debug_prompt_view" not in st.session_state:
            st.session_state.debug_prompt_view = ""
        
        st.text_area(
            "Full prompt sent to Vox-LM",
            key="debug_prompt_view",
            height=400,
            disabled=True)
        
        st.caption(
            "Note: Editing this box does NOT change the prompt sent to the model. It is for debugging and inspection only.")

        if st.session_state.get("handwriting_transcription_debug", ""):
            st.text_area(
                "Prompt sent for handwriting transcription",
                value=st.session_state.handwriting_transcription_debug,
                height=250,
                disabled=True,
                key="handwriting_debug_prompt_view")

#Summarization module in second tab
with tab_summary:
    st.subheader(":violet[Class Performance Summary for Teachers]")
    st.write(
        "*Upload the rubrics/model-answer-graded batch CSV exported from Vox-LM SAQ Marking.* "
        "**_:red[Do not upload the norm-referenced CSV here.]_** "
        "*This summarization module will then analyze the distribution of student scores, identify common strengths and weaknesses, and provide* "
        "*summaries of class performance, misconceptions, and next steps for teaching enhancement.*"
    )

    st.caption(
        "For the best summary, keep the same question, model answer, rubric, and subquestions loaded in the "
        "**:red[SAQ marking tab]**"
    )

    summary_csv = st.file_uploader(
        "Upload rubrics/model-answer-graded CSV (example: voxlm_batch_model_answer_results.csv)",
        type=["csv"],
        key="summary_csv_uploader"
    )

    max_examples_per_tier = st.number_input(
    "Number of representative examples per tier to be included for summarization (maximum of 50)",
    min_value=1,
    max_value=50,
    value=10,
    step=1,
    )

    if st.button("**:blue[Generate summary]**", key="btn_generate_class_summary"):
        if summary_csv is None:
            st.error("You have not uploaded a CSV file. Please upload one to begin class-based summarization.")
        else:
            try:
                try:
                    summary_df = pd.read_csv(summary_csv, dtype_backend="numpy_nullable")
                except TypeError:
                    summary_df = pd.read_csv(summary_csv)
                
                if "result_type" in summary_df.columns:
                    unique_types = set(summary_df["result_type"].dropna().astype(str).str.strip().unique())
                    if "norm_referenced" in unique_types and "rubrics_model_answer_referenced" not in unique_types:
                        st.error("Please upload the rubrics/model-answer-graded CSV, not the norm-referenced CSV.")
                        st.stop()

                if "total_score" not in summary_df.columns:
                    st.error(
                        "This CSV does not appear to be a graded batch output from Vox-LM SAQ Marking. "
                        "It must contain a 'total_score' column."
                    )
                else:
                    csv_text = summary_df.to_csv(index=False)

                    summary_question = {
                        "exam_id": "EXAM",
                        "question_id": "Q1",
                        "stem": question_stem,
                        "max_score": max_score,
                        "subquestions": parsed_subquestions if has_subquestions else [],
                        "model_answer": model_answer,
                        "rubric": global_rubric,
                    }

                    if not any([
                        str(question_stem).strip(),
                        str(model_answer).strip(),
                        str(global_rubric).strip(),
                        bool(parsed_subquestions if has_subquestions else []),
                    ]):
                        st.warning(
                            "Question/model answer/rubric fields appear empty. "
                            "Numerical statistics will still be correct, but rubric-based summarization will be incoherent."
                        )

                    summary_payload = {
                        "csv_text": csv_text,
                        "question": summary_question,
                        "has_subquestions": has_subquestions,
                        "discipline": discipline,
                        "max_examples_per_tier": int(max_examples_per_tier),
                    }

                    res = requests.post(
                        BACKEND_SUMMARY_URL,
                        json=summary_payload,
                        headers={"x-api-key": BACKEND_API_KEY},
                        timeout=300
                    )

                    if res.status_code != 200:
                        st.error(f"Summary backend error: {res.status_code} {res.text}")
                    else:
                        data = res.json()
                        st.session_state.batch_summary_result = data
                        st.session_state.batch_summary_debug_prompt = data.get("debug_prompt", "")
                        st.success("Class summary generated successfully.")

            except Exception as e:
                st.error(f"Failed to generate class summary: {e}")

    summary_result = st.session_state.batch_summary_result

    if summary_result is None:
        st.info("Upload a CSV file and click 'Generate Summary' to view results.")
    else:
        c1, c2, c3, c4, c5, c6 = st.columns(6)

        with c1:
            st.metric("Total Students", summary_result.get("total_students", 0))
        with c2:
            st.metric("Scored by Vox-LM", summary_result.get("scored_students", 0))
        with c3:
            st.metric("Class Average", f"{float(summary_result.get('class_average', 0.0)):.2f}")
        with c4:
            st.metric("Median score", f"{float(summary_result.get('median_score', 0.0)):.2f}")
        with c5:
            st.metric("Maximum score", f"{float(summary_result.get('max_score', 0.0)):.2f}")
        with c6:
            st.metric("Minimum score", f"{float(summary_result.get('min_score', 0.0)):.2f}")

        st.markdown("#### Performance tiers")
        tier_counts = summary_result.get("tier_counts", {}) or {}
        t1, t2, t3 = st.columns(3)
        with t1:
            st.metric("High", int(tier_counts.get("high", 0)))
        with t2:
            st.metric("Intermediate", int(tier_counts.get("mid", 0)))
        with t3:
            st.metric("Low", int(tier_counts.get("low", 0)))

        #st.markdown("### Score distribution")
        #st.json(summary_result.get("score_distribution", {}))

        st.markdown("### Score distribution")

        raw_score_distribution = summary_result.get("score_distribution", {})

        if isinstance(raw_score_distribution, str):
            try:
                score_distribution = json.loads(raw_score_distribution)
            except Exception:
                score_distribution = {}
        elif isinstance(raw_score_distribution, dict):
            score_distribution = raw_score_distribution
        else:
            score_distribution = {}

        preferred_order = [
            "0-24%",
            "25-49%",
            "50-74%",
            "75-100%",
            "lower_third",
            "middle_third",
            "upper_third",
        ]

        if score_distribution:
            rows = []
            for key in preferred_order:
                if key in score_distribution:
                    try:
                        count = int(score_distribution[key])
                    except Exception:
                        count = 0
                    rows.append({"Score interval": str(key), "Number of Students": count})

            for key, val in score_distribution.items():
                if key not in preferred_order:
                    try:
                        count = int(val)
                    except Exception:
                        count = 0
                    rows.append({"Score interval": str(key), "Number of Students": count})

            score_dist_df = pd.DataFrame(rows)
            score_dist_df = sanitize_df_for_streamlit(score_dist_df)

            st.dataframe(score_dist_df, use_container_width=True, hide_index=True)
        else:
            st.info("No score distribution available.")


        overall_sub_stats = summary_result.get("overall_subquestion_stats", {}) or {}
        if overall_sub_stats:
            st.markdown("#### Overall subquestion performance")
            sub_rows = []
            for k, v in overall_sub_stats.items():
                sub_rows.append({
                    "subquestion": k,
                    "mean_score": v.get("mean_score"),
                    "mean_percent_of_sub_max": v.get("mean_percent_of_sub_max"),
                })
            if sub_rows:
                st.dataframe(pd.DataFrame(sub_rows), use_container_width=True)

        st.markdown("#### Overall Summary")
        st.write(summary_result.get("narrative_summary", ""))

        st.markdown("#### Tier-based interpretation")
        tier_summaries = summary_result.get("tier_summaries", {}) or {}
        st.write(f"**High performers:** {tier_summaries.get('high', '')}")
        st.write(f"**Intermediate performers:** {tier_summaries.get('mid', '')}")
        st.write(f"**Low performers:** {tier_summaries.get('low', '')}")

        st.markdown("#### Strengths")
        strengths = summary_result.get("strengths", []) or []
        if strengths:
            for item in strengths:
                st.write(f"- {item}")
        else:
            st.write("No clear strengths identified.")

        st.markdown("#### Weak Areas")
        weak_areas = summary_result.get("weak_areas", []) or []
        if weak_areas:
            for item in weak_areas:
                st.write(f"- {item}")
        else:
            st.write("No major weak areas identified.")

        misconceptions = summary_result.get("common_misconceptions", []) or []
        render_pattern_table(
            "Common misconceptions",
            misconceptions,
            "No common misconceptions identified.",
        )

        out_of_scope_points = summary_result.get("out_of_scope_points", []) or []
        render_pattern_table(
            "Out of scope",
            out_of_scope_points,
            "No major out-of-scope responses identified.")

        st.markdown("#### Common errors by subquestion")
        subquestion_diagnostics = summary_result.get("subquestion_diagnostics", {}) or {}

        if subquestion_diagnostics:
            for sid, info in subquestion_diagnostics.items():
                with st.expander(f"Subquestion {sid}", expanded=False):
                    common_errors = info.get("common_errors", []) or []
                    teaching_note = info.get("teaching_note", "") or ""

                    st.markdown("**Common errors**")
                    if common_errors:
                        for item in common_errors:
                            st.write(f"- {item}")
                    else:
                        st.write("No common errors identified.")

                    st.markdown("**Teaching note**")
                    st.write(teaching_note or "No teaching note returned.")
        else:
            st.info("No subquestion-level diagnostics available.")


        st.markdown("#### Advice to teachers / next steps")
        next_steps = summary_result.get("teacher_next_steps", []) or []
        if next_steps:
            for item in next_steps:
                st.write(f"- {item}")
        else:
            st.write("No teacher advice returned.")

        summary_export = {
            "total_students": summary_result.get("total_students", 0),
            "scored_students": summary_result.get("scored_students", 0),
            "class_average": summary_result.get("class_average", 0.0),
            "median_score": summary_result.get("median_score", 0.0),
            "max_score": summary_result.get("max_score", 0.0),
            "min_score": summary_result.get("min_score", 0.0),
            "std_score": summary_result.get("std_score", 0.0),
            "tier_thresholds": json.dumps(summary_result.get("tier_thresholds", {})),
            "tier_counts": json.dumps(summary_result.get("tier_counts", {})),
            "score_distribution": json.dumps(summary_result.get("score_distribution", {})),
            "overall_subquestion_stats": json.dumps(summary_result.get("overall_subquestion_stats", {})),
            "high_tier_summary": (summary_result.get("tier_summaries", {}) or {}).get("high", ""),
            "mid_tier_summary": (summary_result.get("tier_summaries", {}) or {}).get("mid", ""),
            "low_tier_summary": (summary_result.get("tier_summaries", {}) or {}).get("low", ""),
            "strengths": " | ".join(summary_result.get("strengths", []) or []),
            "common_misconceptions": json.dumps(summary_result.get("common_misconceptions", []) or [], ensure_ascii=False),
            "out_of_scope_points": json.dumps(summary_result.get("out_of_scope_points", []) or [], ensure_ascii=False),
            "weak_areas": " | ".join(summary_result.get("weak_areas", []) or []),
            "teacher_next_steps": " | ".join(summary_result.get("teacher_next_steps", []) or []),
            "narrative_summary": summary_result.get("narrative_summary", ""),
            "subquestion_diagnostics": json.dumps(summary_result.get("subquestion_diagnostics", {})),
        }

        summary_export_df = pd.DataFrame([summary_export])

        summary_csv_buffer = io.StringIO()
        summary_export_df.to_csv(summary_csv_buffer, index=False)
        summary_csv_data = summary_csv_buffer.getvalue()

        st.download_button(
            "Download class summary as CSV",
            data=summary_csv_data,
            file_name="voxlm_class_summary.csv",
            mime="text/csv",
        )

        st.markdown("---")
        st.button(
            "**:red[Reset class performance summary]**",
            key="btn_reset_class_summary",
            on_click=reset_summary_state,
        )

        st.markdown("---")
        st.text_area(
            "Debug prompt sent for class summary",
            value=st.session_state.batch_summary_debug_prompt,
            height=350,
            disabled=True,
            key="summary_debug_prompt_view"
        )

#Student report generation (third tab)
with tab_student_reports:
    st.subheader(":violet[Student Test Reports]")

    st.write(
        "*Upload both rubrics/model-answer-graded CSV and norm-referenced CSV.* "
        "*Vox-LM will generate a concise report for each student.*"
    )

    st.caption(
        "If STUDENT ID is present, it will be used to match and label students. "
        "If not, students will be labelled Student 1, Student 2, etc. using row order in CSV."
    )

    student_report_criterion_csv = st.file_uploader(
        "Upload rubrics/model-answer-graded CSV",
        type=["csv"],
        key="student_report_criterion_csv"
    )

    student_report_norm_csv = st.file_uploader(
        "Upload norm-referenced CSV",
        type=["csv"],
        key="student_report_norm_csv"
    )

    max_total_bullets_per_student = st.number_input(
        "Maximum total bullets per student",
        min_value=1,
        max_value=5,
        value=3,
        step=1,
        key="max_total_bullets_per_student"
    )

    if st.button("**:blue[Generate reports]**", key="btn_generate_student_reports"):
        if student_report_criterion_csv is None or student_report_norm_csv is None:
            st.error("Please upload both CSV files.")
        else:
            try:
                try:
                    criterion_df = pd.read_csv(student_report_criterion_csv, dtype_backend="numpy_nullable")
                except TypeError:
                    criterion_df = pd.read_csv(student_report_criterion_csv)

                try:
                    norm_df = pd.read_csv(student_report_norm_csv, dtype_backend="numpy_nullable")
                except TypeError:
                    norm_df = pd.read_csv(student_report_norm_csv)

                question_dict = {
                    "exam_id": "EXAM",
                    "question_id": "Q1",
                    "stem": question_stem,
                    "max_score": max_score,
                    "subquestions": parsed_subquestions if has_subquestions else [],
                    "model_answer": model_answer,
                    "rubric": global_rubric,
                }

                payload = {
                    "criterion_csv_text": criterion_df.to_csv(index=False),
                    "norm_csv_text": norm_df.to_csv(index=False),
                    "question": question_dict,
                    "has_subquestions": has_subquestions,
                    "discipline": discipline,
                    "max_total_bullets_per_student": int(max_total_bullets_per_student),
                }

                res = requests.post(
                    BACKEND_STUDENT_REPORT_URL,
                    json=payload,
                    headers={"x-api-key": BACKEND_API_KEY},
                    timeout=300
                )

                if res.status_code != 200:
                    st.error(f"Student report backend error: {res.status_code} {res.text}")
                else:
                    data = res.json()
                    st.session_state.student_reports_result = data
                    st.session_state.student_reports_debug_prompt = data.get("debug_prompt_sample", "")
                    st.success("Student reports generated successfully.")

            except Exception as e:
                st.error(f"Failed to generate student reports: {e}")

    student_reports_result = st.session_state.student_reports_result

    if student_reports_result is None:
        st.info("Upload both CSV files and click 'Generate reports' to view results.")
    else:
        solo = student_reports_result.get("solo_question_analysis", {}) or {}

        st.markdown("#### SOLO taxonomy analysis per question/subquestion")
        st.write(solo.get("question_level_summary", ""))

        overall_levels = solo.get("overall_solo_levels", []) or []
        if overall_levels:
            st.write(f"**SOLO levels assessed:** {', '.join(overall_levels)}")

        sub_map = solo.get("subquestion_solo_map", {}) or {}
        if sub_map:
            sub_rows = [{"Subquestion": k, "SOLO level": v} for k, v in sub_map.items()]
            st.dataframe(pd.DataFrame(sub_rows), use_container_width=True, hide_index=True)

        st.markdown("#### Student reports")
        st.caption(
            "For each student, an overall summary is provided along with 2-3 key strengths and 2-3 key areas for improvement. "
            "These are based on both the student's individual performance and how they performed relative to their peers. "
            "Use the download buttons to get a DOCX or PDF report for each student, or download all reports together."
        )

        reports = student_reports_result.get("reports", []) or []
        export_rows = []

        for rep in reports:
            rep = normalize_report_item(rep)
            display_name = rep["display_name"]
            strong = rep["strong_areas"]
            weak = rep["weak_areas"]
            fname = safe_filename(display_name)

            with st.expander(display_name, expanded=False):
                st.write(rep["overall_summary"] or "No summary available.")

                st.markdown("**Strong areas**")
                if strong:
                    for item in strong:
                        st.write(f"- {item}")
                else:
                    st.write("- No major strong areas highlighted.")

                st.markdown("**Weak areas**")
                if weak:
                    for item in weak:
                        st.write(f"- {item}")
                else:
                    st.write("- No major weak areas highlighted.")

                # Build downloadable files for this student
                rep_docx = build_single_report_docx_bytes(rep)
                rep_pdf = build_single_report_pdf_bytes(rep)

                c1, c2 = st.columns(2)
                with c1:
                    st.download_button(
                        "Download DOCX",
                        data=rep_docx,
                        file_name=f"{fname}_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"docx_{fname}",
                    )
                with c2:
                    st.download_button(
                        "Download PDF",
                        data=rep_pdf,
                        file_name=f"{fname}_report.pdf",
                        mime="application/pdf",
                        key=f"pdf_{fname}",
                    )

            export_rows.append({
                "student_id": rep.get("student_id", ""),
                "display_name": display_name,
                "overall_summary": rep.get("overall_summary", ""),
                "strong_areas": " | ".join(strong),
                "weak_areas": " | ".join(weak),
            })

        if export_rows:
            export_df = pd.DataFrame(export_rows)
            export_buffer = io.StringIO()
            export_df.to_csv(export_buffer, index=False)

            st.download_button(
                "Download reports CSV",
                data=export_buffer.getvalue(),
                file_name="voxlm_student_reports.csv",
                mime="text/csv",
            )

        # Download ALL reports as one DOCX or one PDF
        if reports:
            st.markdown("#### Download all reports")

            all_docx = build_all_reports_docx_bytes(reports, solo)
            all_pdf = build_all_reports_pdf_bytes(reports, solo)

            c1, c2 = st.columns(2)
            with c1:
                st.download_button(
                    "**:blue[Download all reports (DOC)]**",
                    data=all_docx,
                    file_name="voxlm_student_reports.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_all_reports_docx",
                )
            with c2:
                st.download_button(
                    "**:blue[Download all reports (PDF)]**",
                    data=all_pdf,
                    file_name="voxlm_student_reports.pdf",
                    mime="application/pdf",
                    key="download_all_reports_pdf",
                )


        st.markdown("---")
        st.markdown("---")
        st.button(
            "**:red[Reset student reports]**",
            key="btn_reset_student_reports",
            on_click=reset_student_reports_state,)

        st.text_area(
            "Debug prompt sent for student reports",
            value=st.session_state.student_reports_debug_prompt,
            height=350,
            disabled=True,
            key="student_reports_debug_prompt_view")
        
        reports = student_reports_result.get("reports", []) or []
        solo = student_reports_result.get("solo_question_analysis", {}) or {}

#mcq from video lectures - tab 4
with tab_mcq_from_videos:
    st.subheader(":violet[Question Generation from Teaching Videos]")
    st.write(
        "Upload a teaching video. Vox-LM will transcribe the audio, analyse selected video frames, "
        "and generate a pre-question plus timestamped in-video questions for teachers to review."
    )

    st.caption(
        "Generates only timestamped educational questions for teachers to review and export. "
        "It does not yet provide an interactive video player."
    )

    video_file = st.file_uploader(
        "Upload video",
        type=["mp4", "mov", "m4v", "avi", "webm"],
        key="video_mcq_video_upload",
    )

    transcript_vtt_file = st.file_uploader(
        "Upload transcript file (optional, .vtt format)",
        type=["vtt"],
        key="video_mcq_vtt_upload",
    )
        
    num_pre_questions = st.number_input(
            "Number of prequestions to generate",
            min_value=1,
            max_value=10,
            value=1,
            step=1,
            key="video_mcq_num_pre_questions",
        )

    question_format = st.selectbox(
            "Question format",
            options=["MCQ", "SAQ", "Mixed"],
            index=0,
            key="video_mcq_question_format",
        )

    c1, c2 = st.columns(2)

    with c1:
        video_topic = st.text_input(
            "Topic / title",
            value="",
            placeholder="Example for dentistry: Management of RPD framework distortion",
            key="video_mcq_topic",
        )

        video_target_level = st.text_input(
            "Target learner level",
            value="",
            placeholder="Example for dentistry: Year 3 dental students",
            key="video_mcq_target_level",
        )

        num_check_questions = st.number_input(
            "Number of questions to generate",
            min_value=1,
            max_value=10,
            value=4,
            step=1,
            key="video_mcq_num_questions",
        )

    with c2:
        allow_anticipatory = st.checkbox(
            "Allow anticipatory questions",
            value=True,
            key="video_mcq_allow_anticipatory",
        )

        use_frame_analysis = st.checkbox(
            "Analyse video frames",
            value=True,
            key="video_mcq_use_frame_analysis",
        )

        frame_interval_seconds = st.number_input(
            "Frame sampling interval, seconds",
            min_value=5,
            max_value=120,
            value=30,
            step=5,
            key="video_mcq_frame_interval",
        )

        max_frames = st.number_input(
            "Maximum frames to analyse",
            min_value=1,
            max_value=100,
            value=20,
            step=1,
            key="video_mcq_max_frames",
        )

    learning_objectives_text = st.text_area(
        "Learning objectives (optional). Please enter one per line.",
        height=120,
        placeholder="Example for dentistry:\nExplain why framework distortion matters\nDifferentiate repairable and non-repairable RPD problems",
        key="video_mcq_learning_objectives",
    )

    generate_video_mcq_btn = st.button(
        "**:blue[Generate Questions]**",
        key="btn_generate_video_mcq",
    )

    if generate_video_mcq_btn:
        if video_file is None:
            st.error("Please upload a video first.")
        else:
            try:
                learning_objectives = [
                    line.strip()
                    for line in learning_objectives_text.splitlines()
                    if line.strip()
                ]

                files = {
                    "video": (
                        video_file.name,
                        video_file.getvalue(),
                        video_file.type or "video/mp4")
                }

                if transcript_vtt_file is not None:
                    files["transcript_vtt"] = (
                        transcript_vtt_file.name,
                        transcript_vtt_file.getvalue(),
                        "text/vtt")

                data = {
                    "discipline": discipline,
                    "topic": video_topic,
                    "target_level": video_target_level,
                    "learning_objectives_json": json.dumps(learning_objectives),
                    "num_pre_questions": str(int(num_pre_questions)),
                    "num_check_questions": str(int(num_check_questions)),
                    "question_format": question_format.lower(),
                    "allow_anticipatory": str(bool(allow_anticipatory)).lower(),
                    "use_frame_analysis": str(bool(use_frame_analysis)).lower(),
                    "frame_interval_seconds": str(int(frame_interval_seconds)),
                    "max_frames": str(int(max_frames)),
                }


                with st.spinner(
                    "Processing video. May take several minutes for longer videos..."
                ):
                    res = requests.post(
                        BACKEND_VIDEO_MCQ_URL,
                        files=files,
                        data=data,
                        headers={"x-api-key": BACKEND_API_KEY},
                        timeout=1200,
                    )

                if res.status_code != 200:
                    st.error(f"Video MCQ backend error: {res.status_code} {res.text}")
                else:
                    result = res.json()
                    st.session_state.video_mcq_session_counter += 1
                    display_id = f"video_{st.session_state.video_mcq_session_counter:03d}"
                    result["display_video_id"] = display_id
                    st.session_state.video_mcq_result = result
                    st.session_state.video_mcq_debug_prompt = result.get("debug_prompt", "")
                    st.session_state.video_mcq_display_id = display_id
                    st.success("Questions generated successfully.")


            except Exception as e:
                st.error(f"Failed to generate questions from video: {e}")

    result = st.session_state.video_mcq_result

    if result is None:
        st.info("Upload a video and click Generate Questions.")
    else:
        st.markdown("---")
        st.markdown("### :violet[Video analysis summary]")

        c1, c2, c3 = st.columns(3)

        with c1:
            st.metric("Session video", result.get("display_video_id", result.get("video_id", "")))

        with c2:
            duration_seconds = result.get("duration_seconds", 0)
            st.metric("Duration", format_seconds(duration_seconds))

        with c3:
            pre_count = len(result.get("pre_questions", []) or [])
            embedded_count = len(result.get("embedded_questions", []) or [])

            st.metric("Generated questions", pre_count + embedded_count)
            st.caption(f"{pre_count} prequestion(s), {embedded_count} in-video question(s)") #may remove

        warnings = result.get("warnings", []) or []
        if warnings:
            st.warning("Warnings")
            for w in warnings:
                st.write(f"- {w}")

        st.markdown("#### :violet[Transcript summary]")
        st.write(result.get("transcript_summary", "") or "No transcript summary returned.")

        st.markdown("#### :violet[Teaching segments]")
        teaching_segments = result.get("teaching_segments", []) or []

        if teaching_segments:
            segment_rows = []
            for seg in teaching_segments:
                segment_rows.append(
                    {
                        "Segment": seg.get("segment_id", ""),
                        "Start": format_seconds(seg.get("start_seconds", "")),
                        "End": format_seconds(seg.get("end_seconds", "")),
                        "Summary": seg.get("summary", ""),
                        "Key concepts": " | ".join(seg.get("key_concepts", []) or []),
                        "Question opportunity": seg.get("suitable_question_opportunity", ""),
                        "Suggested type": seg.get("suggested_question_type", ""),
                    }
                )


            st.dataframe(
                sanitize_df_for_streamlit(pd.DataFrame(segment_rows)),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.info("No teaching segments returned.")

        st.markdown("---")
        st.markdown("### :violet[Generated Questions]")

        pre_question = result.get("pre_question", {}) or {}
        embedded_questions = result.get("embedded_questions", []) or []

        def normalise_question_type(qtype: Any) -> str:
            qtype = str(qtype or "").strip().lower()

            if qtype in ["pre_question", "pre-question", "pre"]:
                return "pre_question"

            if qtype in ["anticipatory", "prediction", "prediction_question"]:
                return "anticipatory"

            if qtype in ["embedded_check", "embedded", "check", "embedded_question"]:
                return "embedded_check"

            return qtype

        def render_video_question(q: Dict[str, Any], title: str):
            st.markdown(f"#### {title}")

            qtype_raw = q.get("question_type", "")
            qtype = normalise_question_type(qtype_raw)
            qkind = str(q.get("question_kind", "mcq") or "mcq").lower()

            ts = q.get("timestamp_seconds", "")
            reveal = q.get("reveal_after_seconds", "")
            evidence_start = q.get("evidence_start_seconds", "")
            evidence_end = q.get("evidence_end_seconds", "")

            st.write(f"**Type:** {qtype_raw}")
            st.write(f"**Format:** {qkind.upper()}")

            cognitive_action = q.get("cognitive_action", "")
            if cognitive_action:
                st.write(f"**Cognitive action:** {cognitive_action}")

            if qtype == "pre_question":
                st.write("**Timing:** Before video starts")

            elif qtype == "anticipatory":
                formatted_ts = format_seconds(ts)
                st.write(f"**Question timestamp:** {formatted_ts or ts}")

                if reveal not in ["", None]:
                    try:
                        reveal_point = float(ts) + float(reveal)
                        st.write(f"**Expected reveal point:** {format_seconds(reveal_point)}")
                    except Exception:
                        st.write(f"**Reveal after:** {format_seconds(reveal) or reveal}")

            elif qtype == "embedded_check":
                formatted_ts = format_seconds(ts)
                st.write(f"**Question timestamp:** {formatted_ts or ts}")

                if evidence_start not in ["", None] and evidence_end not in ["", None]:
                    formatted_start = format_seconds(evidence_start)
                    formatted_end = format_seconds(evidence_end)

                    if formatted_start and formatted_end:
                        st.write(f"**Relevant video section:** {formatted_start}–{formatted_end}")
                    else:
                        st.write(f"**Relevant video section:** {evidence_start}–{evidence_end}")

            else:
                formatted_ts = format_seconds(ts)
                st.write(f"**Timestamp:** {formatted_ts or ts}")

            st.write(f"**Stem:** {q.get('stem', '')}")

            if qkind == "mcq":
                options = q.get("options", []) or []
                for opt in options:
                    if isinstance(opt, dict):
                        st.write(f"**{opt.get('label', '')}.** {opt.get('text', '')}")

                st.write(f"**Correct option:** {q.get('correct_option', '')}")

            else:
                expected_answer = q.get("expected_answer", "")
                marking_points = q.get("marking_points", []) or []

                with st.expander("Expected answer / marking guide"):
                    st.write(expected_answer or "No expected answer returned.")

                    if marking_points:
                        st.markdown("**Marking points**")
                        for mp in marking_points:
                            st.write(f"- {mp}")

            st.write(f"**Difficulty:** {q.get('difficulty', '')}")

            with st.expander("Feedback and teacher rationale"):
                if qkind == "mcq":
                    st.write(f"**Feedback if correct:** {q.get('feedback_correct', '')}")
                    st.write(f"**Feedback if incorrect:** {q.get('feedback_incorrect', '')}")

                st.write(f"**Rationale:** {q.get('rationale', '')}")
                st.write(f"**Learning objective:** {q.get('learning_objective', '')}")
                st.write(f"**Placement reason:** {q.get('placement_reason', '')}")

                qf = q.get("quality_flags", {}) or {}
                if qf:
                    st.write("**Quality flags:**")
                    st.json(qf)

        pre_questions = result.get("pre_questions", []) or []

        # Backward compatibility
        if not pre_questions and result.get("pre_question"):
            pre_questions = [result.get("pre_question")]

        for i, q in enumerate(pre_questions, start=1):
            render_video_question(q, f"Pre-question {i}")

        for i, q in enumerate(embedded_questions, start=1):
            qtype = normalise_question_type(q.get("question_type", ""))

            if qtype == "anticipatory":
                title = f"Anticipatory question {i}"
            elif qtype == "embedded_check":
                title = f"Embedded check question {i}"
            else:
                title = f"Video question {i}"

            render_video_question(q, title)

        st.markdown("---")
        st.markdown("### :violet[Export]")

        export_json = json.dumps(result, indent=2, ensure_ascii=False)

        st.download_button(
            ":green[Download generated question package JSON]",
            data=export_json,
            file_name="voxlm_video_question_package.json",
            mime="application/json",
            key="download_video_mcq_json",
        )

        flat_rows = flatten_video_mcq_rows(result)

        if flat_rows:
            flat_df = pd.DataFrame(flat_rows)
            csv_buffer = io.StringIO()
            flat_df.to_csv(csv_buffer, index=False)

            st.download_button(
                ":green[Download generated questions as CSV]",
                data=csv_buffer.getvalue(),
                file_name="voxlm_video_questions.csv",
                mime="text/csv",
                key="download_video_mcq_csv",
            )

        st.markdown("---")

        if st.button(
            "**:red[Reset video Question generator]**",
            key="btn_reset_video_mcq",
            on_click=reset_video_mcq_state,
        ):
            pass

        with st.expander("Debug prompt sent for Question generation"):
            st.text_area(
                "Question debug prompt",
                value=st.session_state.video_mcq_debug_prompt,
                height=350,
                disabled=True,
                key="video_mcq_debug_prompt_view")
